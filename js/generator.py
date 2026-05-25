import json
from docxtpl import DocxTemplate
import time
import os
import docx
from os.path import join as pj
import pandas as pd
from .utils.docx_stuff import remove_line, remove_page, change_page_to_portrait

 
class GRTask:
    def __init__(self, dictAttributesJson, pathTaskFolder):
        print(f'pathTaskFolder: {pathTaskFolder}')
        self.dfAutomatedSystems = pd.read_excel('xlsxs/indexed_ases.xlsx', engine='openpyxl')
        self.dfProcessesClientPath = pd.read_excel('xlsxs/all_procs.xlsx', engine='openpyxl')

        self.nameAudit = dictAttributesJson['name']
        self.flagPrepOrExecute = dictAttributesJson['pod_or_prov']
        self.dataEmpExecutionControl = dictAttributesJson['director']
        self.dataSignatory = dictAttributesJson['podpis'].split(',')
        self.userLogin = dictAttributesJson['user']
        self.userEmail = dictAttributesJson['email']

        if len(dictAttributesJson['aski']) > 1:
            self.listAutomatedSystemsAudit = [
                col.split('$') for col in dictAttributesJson['aski'].strip(' %').split('%')
            ]
        else:
            self.listAutomatedSystemsAudit = None

        if 'provNum' in dictAttributesJson:
            self.flagIsChange = 1
        else:
            self.flagIsChange = 0

        if dictAttributesJson['isTB'] == "1":
            self.flagIsTB = 1
            self.nameTB = dictAttributesJson['tbName']
        else:
            self.flagIsTB = 0

        self.pathTaskFolder = pathTaskFolder
        self.nameFile = f'Распоряжение_{self.pathTaskFolder}.docx'

        if not self.flagIsChange:
            self.nameTemplate = 'tpl/TemplateOrderExecuteOrPrepAudit2026.docx'
            self.dateStart, self.dateEnd = dictAttributesJson['dates'].split(',')
            self.pointAudit = dictAttributesJson['punkt']
            self.dataHeadAudit = dictAttributesJson['supervisor'].split(',')
            self.listDataEmployeesAudit = [
                col.split(',') for col in dictAttributesJson['emps'].strip(';').split(';')
            ]

            if len(dictAttributesJson['act_date']) > 1:
                self.dateAct = dictAttributesJson['act_date']
            else:
                self.dateAct = None

            if len(dictAttributesJson['subsuper']) > 1:
                self.dataDeputyHeadAudit = dictAttributesJson['subsuper'].split(',')
            else:
                self.dataDeputyHeadAudit = None

            self.listDataCuratorsAudit = [
                cur.split(',') for cur in dictAttributesJson['curator'].split(';')
            ]
            if self.listDataCuratorsAudit[0][0] == '-':
                self.listDataCuratorsAudit = None

            if len(dictAttributesJson['procsKPs']) > 1:
                self.listProcessesAudit = [
                    col.split('$') for col in dictAttributesJson['procsKPs'].strip(' %').split('%')
                ]
            else:
                self.listProcessesAudit = None

        else:
            self.nameTemplate = 'tpl/TemplateOrderChangeAudit2026.docx'
            self.dateOldEndAudit, self.dateNewEndAudit = dictAttributesJson['dates'].split(',')
            self.numberOrderAudit = dictAttributesJson['provNum']
            self.textReason = dictAttributesJson['reason']

            if len(dictAttributesJson['procsKPsAdd']) > 1:
                self.newListProcessesAudit = [
                    col.split('$') for col in dictAttributesJson['procsKPsAdd'].strip(' %').split('%')
                ]
            else:
                self.newListProcessesAudit = None

            if len(dictAttributesJson['procsKPsRemove']) > 1:
                self.deleteListProcessesAudit = [
                    col.split('$') for col in dictAttributesJson['procsKPsRemove'].strip(' %').split('%')
                ]
            else:
                self.deleteListProcessesAudit = None

            if len(dictAttributesJson['addEmployees']) > 1:
                self.listDataEmployeesAdd = [
                    col.split(',') for col in dictAttributesJson['addEmployees'].strip(';').split(';')
                ]
            else:
                self.listDataEmployeesAdd = None

            if len(dictAttributesJson['delEmployees']) > 1:
                self.listDataEmployeesDel = [
                    col.split(',') for col in dictAttributesJson['delEmployees'].strip(';').split(';')
                ]
            else:
                self.listDataEmployeesDel = None

        self.dictNameTB = {
            "ТБ ББ": ["Байкальский банк", "Байкальскому банку"],
            "ТБ ВВБ": ["Волго-Вятский банк", "Волго-Вятскому банку"],
            "ТБ ДВБ": ["Дальневосточный банк", "Дальневосточному банку"],
            "ТБ МБ": ["Московский банк", "Московскому банку"],
            "ТБ ПБ": ["Поволжский банк", "Поволжскому банку"],
            "ТБ СЗБ": ["Северо-Западный банк", "Северо-Западному банку"],
            "ТБ СИБ": ["Сибирский банк", "Сибирскому банку"],
            "ТБ СРБ": ["Среднерусский банк", "Среднерусскому банку"],
            "ТБ УРАЛ": ["Уральский банк", "Уральскому банку"],
            "ТБ ЦЧБ": ["Центрально-Чернозёмный банк", "Центрально-Чернозёмному банку"],
            "ТБ ЮЗБ": ["Юго-Западный банк", "Юго-Западному банку"],
            "ЦА": ["Центральный аппарат", "Центральному аппарату"]
        }

        self.dictNamePost = {
            'Аудитор': 1,
            'Аудитор-младший аналитик данных': 1,
            'Аудитор-младший исследователь данных': 1,
            'Ведущ.спец.по аудиту-исследователь данн.': 2,
            'Ведущий аудитор': 3,
            'Ведущий аудитор-аналитик данных': 3,
            'Ведущий аудитор-исследователь данных': 3,
            'Ведущий специалист по цифр.техн-м аудита': 2,
            'Ведущий эксперт по аудиту': 4,
            'Главный аудитор': 4,
            'Главный аудитор-аналитик данных': 4,
            'Главный аудитор-исследователь данных': 4,
            'Главный инженер по разработке': 4,
            'Главный эксперт по цифр.техн-м аудита': 4,
            'Директор управления': 8,
            'Заместитель директора управления': 7,
            'Заместитель начальника отдела': 5,
            'Исполнительный директор-начальник отдела': 10,
            'Менеджер направления': 5,
            'Менеджер по работе с регионами': 5,
            'Начальник отдела': 6,
            'Рук-ль направления по исслед.данных': 9,
            'Руководитель направления': 9,
            'Специалист': 1,
            'Специалист по цифр.техн-м аудита': 1,
            'Старший аудитор': 2,
            'Старший аудитор-аналитик данных': 2,
            'Старший аудитор-исследователь данных': 2,
            'Старший управляющий директор-дир.управл.': 11,
            'Управляющий директор': 7,
            'Управляющий директор-зам.дир.управления': 7,
            'Эксперт по контролю операций': 3,
            'Эксперт по цифр.техн-м аудита': 3}

        if self.flagPrepOrExecute == '1':
            self.listWordForms = [
                'О проведении аудиторской проверки',
                'к проведению аудиторской проверки',
                'реализацию проверки',
                'проведения проверки',
                'аудиторской проверки ',
                'аудиторской проверки',
                'для проведения аудиторской проверки',
                'проведения проверки',
                'проведение проверки',
                'проверочные действия',
                'приступить',
                'реализацию проверки'
            ]
        else:
            self.listWordForms = [
                'О подготовке к аудиторской проверке',
                'подготовку к аудиторской проверке',
                'подготовку',
                'подготовки к проверки',
                '',
                'подготовки к аудиторской проверке',
                'для подготовки аудиторской проверки',
                'подготовки к аудиторской проверке',
                'подготовку аудиторской проверки',
                '',
                'начать',
                'подготовку к аудиторской проверке'
            ]            
            
    def generate_docx(self):
        """
        Генерирует документ в формате .docx на основе текущего состояния объекта.

        Если атрибут `isChange` установлен в True, вызывается метод `generate_change()`.
        Для генерации распоряжение на изменения.
        В противном случае вызывается метод `generate_prov()`.
        Для генерации распоряжение на подготовку или проведение проверки.

        Returns:
            None: Метод не возвращает значение напрямую, но создает или изменяет документ.
        """
        if self.flagIsChange:
            self.generateChangeAudit()
        else:
            self.generatePrepOrExecuteAudit()


    def generatePrepOrExecuteAudit(self):
        
        # Для сотрудника, исполняющего обязанности контроля исполнения
        empExecutionControlFio, empExecutionControlPost = self.dataEmpExecutionControl.split(',')
        empExecutionControlShortFio = self.short_fio(empExecutionControlFio, accus=True)
        empExecutionControlPost = self.title_to_gent(empExecutionControlPost.replace('управления', ''))
        if self.flagIsTB:
            empExecutionControl = empExecutionControlPost + ' Управления внутреннего аудита по ' + self.dictNameTB[self.nameTB][1] + ' ' + empExecutionControlShortFio
        else:
            empExecutionControl = empExecutionControlPost + ' Управления внутреннего аудита ' + empExecutionControlShortFio

        # Список словарей для генерации таблицы сотрудников
        listDictEmployees = []
        self.listDataEmployeesAudit.sort(key=lambda x: (x[3], self.dictNamePost.get(x[2])), reverse=True)
        for i, emp in enumerate(self.listDataEmployeesAudit):
            dictEmployees = dict()
            dictEmployees['title'] = self.deabbr_title(emp[2])
            dictEmployees['dep'] = emp[1]
            dictEmployees['tb'] = self.dictNameTB[emp[3]][0]
            fio = self.remove_brackets(emp[0])
            dictEmployees['i'] = str(i + 1)
            dictEmployees['fio'] = fio
            dictEmployees['tn'] = emp[0][emp[0].index('(') + 1:-1]
            listDictEmployees.append(dictEmployees)

        # Словари для генерации таблицы процессов
        listDictProcesses = []
        uniqueClientPath = set()
        uniqueProcess = set()
        i_CP = 1
        i_P = 1

        if self.listProcessesAudit:
            # Разбиваем по парам и сортируем для группировки
            listPairsProcessesClientPath = []
            for row in self.listProcessesAudit:
                if '^' in row[1]:
                    CPS = row[1].split('^')
                    for CP in CPS:
                        listPairsProcessesClientPath.append([row[0], CP])
                else:
                    listPairsProcessesClientPath.append(row)
            listPairsProcessesClientPath.sort(key=lambda x: x[1], reverse=True)

            prevCP = 0
            for row in listPairsProcessesClientPath:
                rowProcess = self.dfProcessesClientPath[self.dfProcessesClientPath['Код процесса'] == ('П' + row[0])]
                dictProcess = dict()
                dictProcess['p_name'] = rowProcess['Наименование процесса'].iloc[0]
                dictProcess['p_code'] = rowProcess['Код процесса'].iloc[0]
                dictProcess['p_owner'] = rowProcess['Подразделение - владелец процесса'].iloc[0]
                dictProcess['p_i'] = str(i_P)
                i_P += 1
                # Для подсчёта процессов
                uniqueProcess.add(dictProcess['p_code'])

                # Если нет клиентских путей
                if len(row[1]) <= 1:
                    dictProcess['kp_i'] = '-'
                    dictProcess['kp_code'] = '-'
                    dictProcess['kp_name'] = '-'
                    dictProcess['kp_owner'] = '-'
                    listDictProcesses.append(dictProcess)
                else:
                    codeCP = row[1]
                    if prevCP != codeCP:
                        rowClientPath = rowProcess[rowProcess['Код КП'] == ('КП' + codeCP)]
                        dictProcess = dict(dictProcess)
                        dictProcess['kp_i'] = str(i_CP)
                        dictProcess['kp_name'] = rowClientPath['Наименование КП'].iloc[0]
                        dictProcess['kp_code'] = 'КП' + codeCP
                        dictProcess['kp_owner'] = rowClientPath['Подразделение - владелец КП'].iloc[0]
                        # Для подсчёта КП
                        uniqueClientPath.add(dictProcess['kp_code'])
                        i_CP += 1
                    else:
                        dictProcess['kp_i'] = ''
                        dictProcess['kp_name'] = ''
                        dictProcess['kp_code'] = ''
                        dictProcess['kp_owner'] = ''
                    listDictProcesses.append(dictProcess)
                    prevCP = codeCP

            # Количество процессов для первой страницы
            countProcess = len(uniqueProcess)
            if str(countProcess)[-1] == '1' and countProcess != 11:
                strCountProcess = str(countProcess) + ' централизованный процесс'
            elif str(countProcess)[-1] in ['2', '3', '4'] and countProcess not in [12, 13, 14]:
                strCountProcess = str(countProcess) + ' централизованных процесса'
            else:
                strCountProcess = str(countProcess) + ' централизованных процессов'

            # Сколько чего включить
            countProcessClientPath = ''
            countProcessClientPath += strCountProcess

        # Словари для генерации таблицы с доступами
        if self.listAutomatedSystemsAudit:
            listDictAutomatedSystems = []
            i = 0
            for AS in self.listAutomatedSystemsAudit:
                i += 1
                first = True
                emp_tns = AS[-1][1:-1].split(',')
                rowsAS = self.dfAutomatedSystems[self.dfAutomatedSystems['as_index'] == abs(int(AS[0]))]
                for emp_tn in emp_tns:
                    dictAutomatedSystems = dict()
                    if not first:
                        dictAutomatedSystems['i'] = ''
                        dictAutomatedSystems['name'] = ''
                        dictAutomatedSystems['ke'] = ''
                        dictAutomatedSystems['roles'] = ''
                        dictAutomatedSystems['kp'] = ''
                        dictAutomatedSystems['proc'] = ''
                    else:
                        dictAutomatedSystems['i'] = str(i)
                        dictAutomatedSystems['name'] = rowsAS['title'].iloc[0]
                        dictAutomatedSystems['ke'] = '\r\n'.join([rowsAS['affected_item'].iloc[0], rowsAS['logical_name'].iloc[0]])
                        dictAutomatedSystems['kp'] = AS[-2]
                        dictAutomatedSystems['proc'] = AS[-3]

                        if int(AS[0]) < 0:
                            dictAutomatedSystems['name'] = dictAutomatedSystems['name'] + ' (+реплика)'

                        if AS[1] == '-1':
                            dictAutomatedSystems['roles'] = 'Роль, позволяющая просматривать необходимую информацию в рамках проверки.'
                        else:
                            dictAutomatedSystems['roles'] = '\r\n'.join([rowsAS[rowsAS['role_index'] == int(role_index)]['name'].iloc[0] for role_index in AS[1:-3]])

                    if emp_tn == '-1':
                        dictAutomatedSystems['emp_name'] = 'Все участники (Приложение 1)'
                    else:
                        cur_emp = None
                        for dictEmployees in listDictEmployees:
                            if dictEmployees['tn'] == emp_tn:
                                cur_emp = dictEmployees
                                break
                        if cur_emp:
                            dictAutomatedSystems['emp_name'] = cur_emp['fio']
                            dictAutomatedSystems['title'] = self.get_full_title(cur_emp['title'], cur_emp['tb'])
                            dictAutomatedSystems['tn'] = cur_emp['tn']

                    first = False
                    listDictAutomatedSystems.append(dictAutomatedSystems)

        # Форматирование дат
        if '/' in self.dateStart:
            startDatePart = self.dateStart.split('/')
            startDateFull = startDatePart[1].zfill(2) + '.' + startDatePart[0].zfill(2) + '.' + startDatePart[2]
        else:
            startDateFull = ''

        if '/' in self.dateEnd:
            endDatePart = self.dateEnd.split('/')
            endDateFull = endDatePart[1].zfill(2) + '.' + endDatePart[0].zfill(2) + '.' + endDatePart[2]
        else:
            endDateFull = ''

        if self.dateAct:
            dateActPart = self.dateAct.split('/')
            dateActFull = dateActPart[1].zfill(2) + '.' + dateActPart[0].zfill(2) + '.' + dateActPart[2]
        else:
            dateActFull = ''

        # Куратор
        stringCurators = ''
        plurCurators = 'куратора'
        if self.listDataCuratorsAudit:
            stringCurators = ''
            for i, curator in enumerate(self.listDataCuratorsAudit):
                curatorFIO = self.short_fio(self.remove_brackets(curator[0]), gent=True)
                curatorPost = self.get_full_title(curator[1].lower(), curator[2], abbr=True if i==0 else True, tb=True, gent=True)
                stringCurators += f'{curatorFIO} – {curatorPost}, '
            stringCurators = stringCurators[:-2]
            if len(self.listDataCuratorsAudit) > 1:
                plurCurators = 'кураторов'

        # Руководитель
        headAuditFIO = self.short_fio(self.remove_brackets(self.dataHeadAudit[2]), gent=True)
        headAuditPost = self.get_full_title(self.dataHeadAudit[0].lower(), self.dataHeadAudit[-1], abbr=False, tb=True, gent=True)

        # Зам руководителя
        if self.dataDeputyHeadAudit:
            deputyHeadFIO = self.short_fio(self.remove_brackets(self.dataDeputyHeadAudit[2]), gent=True)
            deputyHeadPost = self.get_full_title(self.dataDeputyHeadAudit[0].lower(), self.dataDeputyHeadAudit[-1], abbr=True, tb=True, gent=True)
        else:
            deputyHeadFIO = deputyHeadPost = ''
        print(self.flagIsTB)
        # Подписант
        if not self.flagIsTB:
            signatoryFIO = self.short_fio(self.dataSignatory[0])
            signatoryPost = self.dataSignatory[1].replace('зам.дир.', 'заместитель директора ') \
                .replace('дир.управл.', 'директор управления') \
                .replace('управления', 'Управления внутреннего аудита')
        else:
            signatoryFIO = self.short_fio(self.dataSignatory[0])
            fullNameTB = self.dictNameTB[self.nameTB][1]
            signatoryPost = self.dataSignatory[1].replace('управления', 'Управления внутреннего аудита') + ' по ' + fullNameTB

        dictOrder = {'name': self.nameAudit.strip(),
                     'pv_title': self.listWordForms[0],
                     'pv1': self.listWordForms[1],
                     'pv2': self.listWordForms[2],
                     'pv3': self.listWordForms[3],
                     'pv4': self.listWordForms[4],
                     'pv5': self.listWordForms[5],
                     'pv9': self.listWordForms[9],
                     'pv10': self.listWordForms[10],
                     'pv_emp': self.listWordForms[6],
                     'pv_ases': self.listWordForms[7],
                     'podpunkt': self.pointAudit,
                        'start_date': startDateFull,
                        'end_date': endDateFull,
                        'act_date': dateActFull,
                        'curator': stringCurators,
                        'supervisor': headAuditFIO + ' – ' + headAuditPost,
                        'subsuper': deputyHeadFIO + ' – ' + deputyHeadPost,
                        'curator_plur': plurCurators,
                        'dir': empExecutionControl,
                        'podpis_title': signatoryPost,
                        'podpis_fio': signatoryFIO,
                        'employees': listDictEmployees,
                        'pril_procs': 2,
                        'pril_aski': 3,
                        }

        temp_tpl_doc = docx.Document(self.nameTemplate)
        for dict_emp in listDictEmployees:
            if 'Центральный аппарат' in dict_emp['tb']:
                ca_tb = "Центральный аппарат/Территориальный банк"
                break
            else:
                ca_tb = "Территориальный банк"
        dictOrder['ca_tb'] = ca_tb    


        if self.listProcessesAudit:
            dictOrder['kps_and_ps_count'] = countProcessClientPath
            dictOrder['processes'] = listDictProcesses
        else:
            dictOrder['pril_aski'] = 2
            remove_page(temp_tpl_doc, 2)

        if self.flagIsTB:
            dictOrder['processes'] = listDictProcesses
            dictOrder['TB_full_name'] = (self.dictNameTB[self.nameTB][0]).upper()
            dictOrder['TB_full_name2'] = 'ПО ' + (self.dictNameTB[self.nameTB][1]).upper()
        else:
            remove_line(temp_tpl_doc, '{{ TB_full_name }}')

        if self.listAutomatedSystemsAudit:
            dictOrder['aski'] = listDictAutomatedSystems
        elif self.listProcessesAudit:
            remove_page(temp_tpl_doc, 3)
        else:
            remove_page(temp_tpl_doc, 2)

        if not self.listProcessesAudit:
            remove_line(temp_tpl_doc, 'В периметр')
        if not self.listAutomatedSystemsAudit:
            remove_line(temp_tpl_doc, 'Предоставить участникам')
        if not self.dateAct:
            remove_line(temp_tpl_doc, 'Направить Акт')
        if not self.dataDeputyHeadAudit:
            remove_line(temp_tpl_doc, 'заместителя руководителя')


        temp_tpl_doc_path = pj('queue', self.userLogin, self.pathTaskFolder, 'temp_tpl.docx')
        temp_tpl_doc.save(temp_tpl_doc_path)

        # Сама генерация
        self.tpl = DocxTemplate(temp_tpl_doc_path)
        self.tpl.render(dictOrder)

        if self.listProcessesAudit:
            # Смержить пустые ячейки в таблице с процессами
            for j in range(4):
                cells = self.tpl.docx.tables[2].columns[j].cells
                for i in range(1, len(cells)):
                    if cells[i].text == '':
                        cells[i].merge(cells[i - 1])
            as_table_num = 3
        else:
            as_table_num = 2

        if self.listAutomatedSystemsAudit:
            # Смержить пустые ячейки в таблице с доступами
            for j in range(4):
                cells = self.tpl.docx.tables[as_table_num].columns[j].cells
                for i in range(1, len(cells)):
                    if cells[i].text == '':
                        cells[i].merge(cells[i - 1])

            # Смержить там где 'Все участники'
            for row in self.tpl.docx.tables[as_table_num].rows:
                cells = row.cells
                if cells[4].text == 'Все участники (Приложение 1)':
                    cells[4].merge(cells[6])

        if not self.listDataCuratorsAudit:
            remove_line(self.tpl.docx, 'куратора проверки')

        os.remove(temp_tpl_doc_path)
        self.save_docx(self.pathTaskFolder)


    # Генерация распоряжения на изменение
    def generateChangeAudit(self):
        # Для директора
        empExecutionControlFio, empExecutionControlPost = self.dataEmpExecutionControl.split(',')
        empExecutionControlShortFio = self.short_fio(empExecutionControlFio, accus=True)
        empExecutionControlPost = self.title_to_gent(empExecutionControlPost.replace('управления', ''))
        if self.flagIsTB:
            empExecutionControl = empExecutionControlPost + ' Управления внутреннего аудита по ' + self.dictNameTB[self.nameTB][1] + ' ' + empExecutionControlShortFio
        else:
            empExecutionControl = empExecutionControlPost + ' Управления внутреннего аудита ' + empExecutionControlShortFio

        # Список словарей для генерации таблицы сотрудников
        listDictEmployees = []
        if self.listDataEmployeesAdd:
            self.listDataEmployeesAdd.sort(key=lambda x: (x[3], self.dictNamePost.get(x[2])), reverse=True)
            for i, emp in enumerate(self.listDataEmployeesAdd):
                dictEmployees = dict()
                dictEmployees['title'] = self.deabbr_title(emp[2])
                dictEmployees['dep'] = emp[1]
                dictEmployees['tb'] = self.dictNameTB[emp[3]][0]
                fio = self.remove_brackets(emp[0])
                dictEmployees['i'] = str(i + 1)
                dictEmployees['fio'] = fio
                dictEmployees['tn'] = emp[0][emp[0].index('(') + 1:-1]
                dictEmployees['action'] = 'Включить'
                listDictEmployees.append(dictEmployees)
        if self.listDataEmployeesDel:
            self.listDataEmployeesDel.sort(key=lambda x: (x[3], self.dictNamePost.get(x[2])), reverse=True)
            for i, emp in enumerate(self.listDataEmployeesDel):
                dictEmployees = dict()
                dictEmployees['title'] = self.deabbr_title(emp[2])
                dictEmployees['dep'] = emp[1]
                dictEmployees['tb'] = self.dictNameTB[emp[3]][0]
                fio = self.remove_brackets(emp[0])
                dictEmployees['i'] = str(i + 1)
                dictEmployees['fio'] = fio
                dictEmployees['tn'] = emp[0][emp[0].index('(') + 1:-1]
                dictEmployees['action'] = 'Исключить'
                listDictEmployees.append(dictEmployees)


        # Словари для генерации таблицы процессов
        listDictProcessesAdd = []
        uniqueClientPath = set()
        uniqueProcess = set()
        i_CP = 1
        i_P = 1
        if self.newListProcessesAudit:
            # Разбиваем по парам и сортируем для группировки
            listPairsProcessesClientPath = []
            for row in self.newListProcessesAudit:
                if '^' in row[1]:
                    CPS = row[1].split('^')
                    for CP in CPS:
                        listPairsProcessesClientPath.append([row[0], CP])
                else:
                    listPairsProcessesClientPath.append(row)
            listPairsProcessesClientPath.sort(key=lambda x: x[1], reverse=True)
            
            prevCP = 0
            for row in listPairsProcessesClientPath:
                rowProcess = self.dfProcessesClientPath[self.dfProcessesClientPath['Код процесса'] == ('П' + row[0])]
                dictProcess = dict()
                dictProcess['p_name'] = rowProcess['Наименование процесса'].iloc[0]
                dictProcess['p_code'] = rowProcess['Код процесса'].iloc[0]
                dictProcess['p_owner'] = rowProcess['Подразделение - владелец процесса'].iloc[0]
                dictProcess['p_i'] = str(i_P)
                i_P += 1
                # Для подсчёта процессов
                uniqueProcess.add(dictProcess['p_code'])

                # Если нет клиентских путей
                if len(row[1]) <= 1:
                    dictProcess['kp_i'] = '-'
                    dictProcess['kp_code'] = '-'
                    dictProcess['kp_name'] = '-'
                    dictProcess['kp_owner'] = '-'
                    listDictProcessesAdd.append(dictProcess)
                else:
                    codeCP = row[1]
                    if prevCP != codeCP:
                        rowClientPath = rowProcess[rowProcess['Код КП'] == ('КП' + codeCP)]
                        dictProcess = dict(dictProcess)
                        dictProcess['kp_i'] = str(i_CP)
                        dictProcess['kp_name'] = rowClientPath['Наименование КП'].iloc[0]
                        dictProcess['kp_code'] = 'КП' + codeCP
                        dictProcess['kp_owner'] = rowClientPath['Подразделение - владелец КП'].iloc[0]
                        # Для подсчёта КП
                        uniqueClientPath.add(dictProcess['kp_code'])
                        i_CP += 1
                    else:
                        dictProcess['kp_i'] = ''
                        dictProcess['kp_name'] = ''
                        dictProcess['kp_code'] = ''
                        dictProcess['kp_owner'] = ''
                    listDictProcessesAdd.append(dictProcess)
                    prevCP = codeCP

            # Количество процессов для первой страницы
            countProcess = len(uniqueProcess)
            if str(countProcess)[-1] == '1' and countProcess != 11:
                strCountProcess = str(countProcess) + ' централизованный процесс'
            elif str(countProcess)[-1] in ['2', '3', '4'] and countProcess not in [12, 13, 14]:
                strCountProcess = str(countProcess) + ' централизованных процесса'
            else:
                strCountProcess = str(countProcess) + ' централизованных процессов'

            # Сколько чего включить
            countProcessClientPathAdd = ''
            countProcessClientPathAdd += strCountProcess
            
        # Словари для генерации таблицы процессов
        listDictProcessesRem = []
        uniqueClientPath = set()
        uniqueProcesses = set()
        i_CP = 1
        i_P = 1
        if self.deleteListProcessesAudit:
            # Разбиваем по парам и сортируем для группировки
            listProcessClientPathPairs = []
            for row in self.deleteListProcessesAudit:
                if '^' in row[1]:
                    CPS = row[1].split('^')
                    for CP in CPS:
                        listProcessClientPathPairs.append([row[0], CP])
                else:
                    listPairsProcessesClientPath.append(row)
            listProcessClientPathPairs.sort(key=lambda x: x[1], reverse=True)

            prevCP = 0
            for row in listProcessClientPathPairs:
                rowProcesses = self.dfProcessesClientPath[self.dfProcessesClientPath['Код процесса'] == ('П' + row[0])]
                dictProcesses = dict()
                dictProcesses['p_name'] = rowProcesses['Наименование процесса'].iloc[0]
                dictProcesses['p_code'] = rowProcesses['Код процесса'].iloc[0]
                dictProcesses['p_owner'] = rowProcesses['Подразделение - владелец процесса'].iloc[0]
                dictProcesses['p_i'] = str(i_P)
                i_P += 1
                # Для подсчёта процессов
                uniqueProcesses.add(dictProcesses['p_code'])

                # Если нет клиентских путей
                if len(row[1]) <= 1:
                    dictProcesses['kp_i'] = '-'
                    dictProcesses['kp_code'] = '-'
                    dictProcesses['kp_name'] = '-'
                    dictProcesses['kp_owner'] = '-'
                    listDictProcessesRem.append(dictProcesses)
                else:
                    codeCP = row[1]
                    if prevCP != codeCP:
                        rowClientPath = rowProcesses[rowProcesses['Код КП'] == ('КП' + codeCP)]
                        dictProcesses = dict(dictProcesses)
                        dictProcesses['kp_i'] = str(i_CP)
                        dictProcesses['kp_name'] = rowClientPath['Наименование КП'].iloc[0]
                        dictProcesses['kp_code'] = 'КП' + codeCP
                        dictProcesses['kp_owner'] = rowClientPath['Подразделение - владелец КП'].iloc[0]
                        # Для подсчёта КП
                        uniqueClientPath.add(dictProcesses['kp_code'])
                        i_CP += 1
                    else:
                        dictProcesses['kp_i'] = ''
                        dictProcesses['kp_name'] = ''
                        dictProcesses['kp_code'] = ''
                        dictProcesses['kp_owner'] = ''
                    listDictProcessesRem.append(dictProcesses)
                    prevCP = codeCP

            # Количество процессов для первой страницы
            countProcess = len(uniqueProcesses)
            if str(countProcess)[-1] == '1' and countProcess != 11:
                strCountProcess = str(countProcess) + ' централизованный процесс'
            elif str(countProcess)[-1] in ['2', '3', '4'] and countProcess not in [12, 13, 14]:
                strCountProcess = str(countProcess) + ' централизованных процесса'
            else:
                strCountProcess = str(countProcess) + ' централизованных процессов'

            # Сколько чего включить
            countRemoveProcessesClientPath = '' 
            countRemoveProcessesClientPath += strCountProcess

        # Словари для генерации таблицы с доступами
        if self.listAutomatedSystemsAudit:
            listDictAutomatedSystems = []
            i = 0
            for AS in self.listAutomatedSystemsAudit:
                i += 1
                first = True
                emps = AS[-1][1:-1].split(',')
                rowsAS = self.dfAutomatedSystems[self.dfAutomatedSystems['as_index'] == abs(int(AS[0]))]
                for emp in emps:
                    dictAutomatedSystems = dict()
                    if not first:
                        dictAutomatedSystems['i'] = ''
                        dictAutomatedSystems['name'] = ''
                        dictAutomatedSystems['ke'] = ''
                        dictAutomatedSystems['roles'] = ''
                        dictAutomatedSystems['kp'] = ''
                        dictAutomatedSystems['proc'] = ''
                    else:
                        dictAutomatedSystems['i'] = str(i)
                        dictAutomatedSystems['name'] = rowsAS['title'].iloc[0]
                        dictAutomatedSystems['ke'] = '\r\n'.join([rowsAS['affected_item'].iloc[0], rowsAS['logical_name'].iloc[0]])
                        dictAutomatedSystems['kp'] = AS[-2]
                        dictAutomatedSystems['proc'] = AS[-3]
                        
                        if int(AS[0]) < 0:
                            dictAutomatedSystems['name'] = dictAutomatedSystems['name'] + ' (+реплика)'

                        if AS[1] == '-1':
                            dictAutomatedSystems['roles'] = 'Роль, позволяющая просматривать необходимую информацию в рамках проверки.'
                        else:
                            dictAutomatedSystems['roles'] = '\r\n'.join([rowsAS[rowsAS['role_index'] == int(role_index)]['name'].iloc[0] for role_index in AS[1:-3]])

                    if emp == '-1':
                        dictAutomatedSystems['emp_name'] = 'Все участники (Приложение 1)'
                    else:
                        emp_tn_fio, emp_title, emp_tb = emp.split('^')
                        emp_fio = self.remove_brackets(emp_tn_fio)
                        emp_tn = emp_tn_fio[emp_tn_fio.index('(') + 1:-1]
                        
                        dictAutomatedSystems['emp_name'] = emp_fio
                        dictAutomatedSystems['title'] = self.get_full_title(emp_title, emp_tb)
                        dictAutomatedSystems['tn'] = emp_tn

                    first = False
                    listDictAutomatedSystems.append(dictAutomatedSystems)

        # Форматирование дат
        if '/' in self.dateOldEndAudit:
            startDatePart = self.dateOldEndAudit.split('/')
            startDateFull = startDatePart[1].zfill(2) + '.' + startDatePart[0].zfill(2) + '.' + startDatePart[2]
        else:
            startDateFull = ''

        if '/' in self.dateNewEndAudit:
            endDatePart = self.dateNewEndAudit.split('/')
            endDateFull = endDatePart[1].zfill(2) + '.' + endDatePart[0].zfill(2) + '.' + endDatePart[2]
        else:
            endDateFull = ''

        if self.dateAct:
            dateActPart = self.dateAct.split('/')
            dateActFull = dateActPart[1].zfill(2) + '.' + dateActPart[0].zfill(2) + '.' + dateActPart[2]
        else:
            dateActFull = ''

        # Подписант
        if not self.flagIsTB:
            signatoryFIO = self.short_fio(self.dataSignatory[0])
            signatoryPost = self.dataSignatory[1].replace('зам.дир.', 'заместитель директора ') \
                .replace('дир.управл.', 'директор управления') \
                .replace('управления', 'Управления внутреннего аудита')
        else:
            signatoryFIO = self.short_fio(self.dataSignatory[0])
            fullNameTB = self.dictNameTB[self.nameTB][1]
            signatoryPost = self.dataSignatory[1].replace('управления', 'Управления внутреннего аудита') + ' по ' + fullNameTB

        # Причина
        if self.textReason:
            stringReason = f'В связи с необходимостью {self.textReason} внести'
        else:
            stringReason = 'Внести'
        
        dictOrder = {'name': self.nameAudit.strip(),
                    'pv_title': f'{self.listWordForms[0]} {self.nameAudit.strip()}',
                    'pv_num': self.numberOrderAudit,
                    'reason': stringReason,
                    'pv5': self.listWordForms[5],
                    'pv8': self.listWordForms[8],
                    'pv11': self.listWordForms[11],
                    'pv_ases': self.listWordForms[7],
                    'pv_date': startDateFull,
                    'end_date': endDateFull,
                    'act_date': dateActFull,
                    'dir': empExecutionControl,
                    'podpis_title': signatoryPost,
                    'podpis_fio': signatoryFIO,
                    'pril_emp': 1,
                    'pril_add': 2,
                    'pril_rem': 3,
                    'pril_aski': 4
                    }

        temp_tpl_doc = docx.Document(self.nameTemplate)

        if self.flagIsTB:
            dictOrder['TB_full_name'] = (self.dictNameTB[self.nameTB][0]).upper()
            dictOrder['TB_full_name2'] = 'ПО ' + (self.dictNameTB[self.nameTB][1]).upper()
        else:
            remove_line(temp_tpl_doc, '{{ TB_full_name }}')

        if self.listDataEmployeesAdd or self.listDataEmployeesDel:
            dictOrder['employees'] = listDictEmployees
            for dict_emp in listDictEmployees:
                if 'Центральный аппарат' in dict_emp['tb']:
                    ca_tb = "Центральный аппарат/Территориальный банк"
                    break
                else:
                    ca_tb = "Территориальный банк"
            dictOrder['ca_tb'] = ca_tb   
        else:
            dictOrder['pril_add'] = 1
            remove_page(temp_tpl_doc, 1)

        if self.newListProcessesAudit:
            dictOrder['kps_and_ps_count_add'] = countProcessClientPathAdd
            dictOrder['processesAdded'] = listDictProcessesAdd
        elif self.listDataEmployeesAdd or self.listDataEmployeesDel:
            dictOrder['pril_rem'] = 2
            remove_page(temp_tpl_doc, 2)
        else:
            dictOrder['pril_rem'] = 1
            remove_page(temp_tpl_doc, 1)

        if self.deleteListProcessesAudit:
            dictOrder['kps_and_ps_count_rem'] = countRemoveProcessesClientPath
            dictOrder['processesRemoved'] = listDictProcessesRem
        elif self.listDataEmployeesAdd or self.listDataEmployeesDel:
            dictOrder['pril_aski'] = 3
            remove_page(temp_tpl_doc, 3)
        elif self.newListProcessesAudit:
            dictOrder['pril_aski'] = 2
            remove_page(temp_tpl_doc, 2)
        else:
            dictOrder['pril_aski'] = 1
            remove_page(temp_tpl_doc, 1)

        if self.listAutomatedSystemsAudit:
            dictOrder['aski'] = listDictAutomatedSystems
        elif (self.listDataEmployeesAdd or self.listDataEmployeesDel) and self.newListProcessesAudit and self.deleteListProcessesAudit:
            remove_page(temp_tpl_doc, 4)
        elif self.newListProcessesAudit and self.deleteListProcessesAudit:
            remove_page(temp_tpl_doc, 3)
        elif (self.listDataEmployeesAdd or self.listDataEmployeesDel) and (self.newListProcessesAudit or self.deleteListProcessesAudit):
            remove_page(temp_tpl_doc, 3)
        elif (self.listDataEmployeesAdd or self.listDataEmployeesDel) or self.newListProcessesAudit or self.deleteListProcessesAudit:
            remove_page(temp_tpl_doc, 2)
        else:
            remove_page(temp_tpl_doc, 1)
            change_page_to_portrait(temp_tpl_doc)

        if not self.dateAct:
            remove_line(temp_tpl_doc, 'Направить Акт')
        if not self.listDataEmployeesAdd and not self.listDataEmployeesDel:
            remove_line(temp_tpl_doc, 'Изменить состав участников в соответствии')
        if not self.newListProcessesAudit:
            remove_line(temp_tpl_doc, 'В периметр')
        if not self.deleteListProcessesAudit:
            remove_line(temp_tpl_doc, 'Из периметра')
        if not self.listAutomatedSystemsAudit:
            remove_line(temp_tpl_doc, 'Предоставить участникам')
        if len(self.dateNewEndAudit) < 2:
            remove_line(temp_tpl_doc, 'Продлить')

        temp_tpl_doc_path = pj('queue', self.userLogin, self.pathTaskFolder, 'temp_tpl.docx')
        temp_tpl_doc.save(temp_tpl_doc_path)

        # Сама генерация
        self.tpl = DocxTemplate(temp_tpl_doc_path)
        self.tpl.render(dictOrder)

        if self.newListProcessesAudit:
            # Смержить пустые ячейки в таблице с процессами
            for j in range(4):
                cells = self.tpl.docx.tables[1].columns[j].cells
                for i in range(1, len(cells)):
                    if cells[i].text == '':
                        cells[i].merge(cells[i - 1])
        if self.deleteListProcessesAudit:
            if self.newListProcessesAudit:
                proc_table_num = 2
            else:
                proc_table_num = 1
            # Смержить пустые ячейки в таблице с процессами
            for j in range(4):
                cells = self.tpl.docx.tables[proc_table_num].columns[j].cells
                for i in range(1, len(cells)):
                    if cells[i].text == '':
                        cells[i].merge(cells[i - 1])
        if self.listAutomatedSystemsAudit:
            # Смержить пустые ячейки в таблице с доступами
            if self.newListProcessesAudit and self.deleteListProcessesAudit:
                as_table_num = 3
            elif self.newListProcessesAudit or self.deleteListProcessesAudit:
                as_table_num = 2
            else:
                as_table_num = 1

            # Смержить пустые ячейки в таблице с доступами
            for j in range(4):
                cells = self.tpl.docx.tables[as_table_num].columns[j].cells
                for i in range(1, len(cells)):
                    if cells[i].text == '':
                        cells[i].merge(cells[i - 1])

            # Смержить там где 'Все участники'
            for row in self.tpl.docx.tables[as_table_num].rows:
                cells = row.cells
                if cells[4].text == 'Все участники (Приложение 1)':
                    cells[4].merge(cells[6])

        os.remove(temp_tpl_doc_path)
        self.save_docx(self.pathTaskFolder)

    def save_docx(self, task_name: str):
        os.makedirs(pj('queue', self.userLogin, task_name), exist_ok=True)
        self.tpl.save(pj('queue', self.userLogin, task_name, self.nameFile))

    def short_fio(self, fio, gent=False, accus=False):
        splited = fio.split()
        if len(splited) == 2:
            if gent:
                surname = self.fio_to_gent(splited[0], '')
            elif accus:
                surname = self.fio_to_accus(splited[0], '')
            else:
                surname = splited[0]
            return splited[1][0] + '. ' + surname
        else:
            if gent:
                surname = self.fio_to_gent(splited[0], splited[2])
            elif accus:
                surname = self.fio_to_accus(splited[0], splited[2]) 
            else:
                surname = splited[0]
            return splited[1][0] + '.' + splited[2][0] + '. ' + surname

    @staticmethod
    def remove_brackets(string):
        return string[:string.rfind('(')].strip()

    def get_full_title(self, title, uprav, tb=True, abbr=True, gent=False):
        full_title = self.deabbr_title(title)
        if gent:
            full_title = self.title_to_gent(full_title)
        if abbr:
            full_title += ' УВА'
        else:
            full_title += ' Управления внутреннего аудита (далее – УВА)'

        if tb:
            if 'ТБ' in uprav:
                full_title += ' по ' + self.dictNameTB[uprav][1]
        return full_title

    def fio_to_gent(self, surname, patron):
        if surname[-1] in 'еёиоуыэю' or surname[-2:] in ['ых', 'их']:
            return surname
        if surname[-2:] == 'ая':
            return surname[:-2] + 'ой'
        if surname[-1] == 'а':
            if surname[-2] in 'вн':
                return surname[:-1] + 'ой'
            if surname[-2] in 'тдз':
                return surname[:-1] + 'ы'
            else:
                return surname[:-1] + 'и'
        if surname[-2:] == 'ия':
            return surname[:-2] + 'ии'
        if surname[-1] == 'я':
            return surname[:-1] + 'и'
        if surname[-2:] == 'ий':
            return surname[:-2] + 'ого'
        if not self.isFemale(patron):
            if surname[-1] not in 'йць':
                return surname + 'а'
        return surname
    
    def fio_to_accus(self, surname, patron):
        if surname[-1] in 'еёиоуыэю' or surname[-2:] in ['ых', 'их']:
            return surname
        if surname[-2:] == 'ая':
            return surname[:-2] + 'ую'
        if surname[-1] == 'а':
            return surname[:-1] + 'у'
        if surname[-2:] == 'ия':
            return surname[:-2] + 'ии'
        if surname[-1] == 'я':
            return surname[:-1] + 'ю'
        if not self.isFemale(patron):
            if surname[-2:] == 'ий':
                return surname[:-2] + 'ого'
            if surname[-1] not in 'йць':
                return surname + 'а'
        return surname

    @staticmethod
    def isFemale(patron):
        if not patron:
            return True
        if patron[-2:] == 'на':
            return True
        return False

    @staticmethod
    def deabbr_title(title):
        return title.replace('Ведущ.', 'Ведущий ') \
            .replace('ведущ.', 'ведущий ') \
            .replace('спец.', 'специалист ') \
            .replace('данн.', 'данных') \
            .replace('цифр.техн-м', 'цифровым технологиям') \
            .replace('рук-ль', 'руководитель') \
            .replace('зам.дир.', 'заместитель директора ') \
            .replace('дир.', 'директор ') \
            .replace('управл.', 'управления') \
            .replace('исслед.', 'исследованию ')

    @staticmethod
    def title_to_gent(title):
        gent_title = []
        for words in title.split():
            tokens = []
            for word in words.split('-'):
                if word.lower() in ['аудитор', 'аналитик', 'специалист', 'эксперт', 'инженер', 'директор', 'менеджер', 'начальник']:
                    tokens.append(word + 'а')
                elif word.lower() in ['исследователь', 'руководитель', 'заместитель']:
                    tokens.append(word[:-1] + 'я')
                elif word.lower() in ['старший', 'управляющий', 'ведущий', 'младший']:
                    tokens.append(word[:-2] + 'его')
                elif word.lower() in ['главный', 'исполнительный']:
                    tokens.append(word[:-2] + 'ого')
                else:
                    tokens.append(word)
            gent_title.append('-'.join(tokens))
        return ' '.join(gent_title)
